import re
import json
import pandas as pd
from datetime import datetime
from pathlib import Path
from typing import List, Union, Any, Dict
import openpyxl # Importar openpyxl para estilos
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Importamos tu modelo (ajusta la ruta según tu estructura)
from app.models.audit import AuditReport

# ReportLab imports
from reportlab.lib import colors
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    PageBreak, Preformatted, Image
)
from reportlab.lib.enums import TA_JUSTIFY, TA_CENTER, TA_LEFT

class ReportGenerator:
    def __init__(self, audit: AuditReport):
        self.audit = audit

        # 1. Datos
        self.lh_data = self.audit.lighthouse_data or {}
        self.seo_data = self.audit.seo_analysis or {}
        self.ai_data = self.audit.ai_suggestions or {}

        # 2. URL y Directorios
        self.url = self.lh_data.get('url') or self.seo_data.get('onpage_seo', {}).get('canonical') or f"audit-{self.audit.id}"

        clean_domain = self.url.replace('https://', '').replace('http://', '').split('/')[0]
        clean_domain = re.sub(r'[^\w\-_\.]', '_', clean_domain)

        self.base_dir = Path("storage/reports") / clean_domain
        self.base_dir.mkdir(parents=True, exist_ok=True)
        self.timestamp = datetime.now().strftime("%Y%m%d_%H%M")

        # 3. Estilos
        self._setup_pdf_styles()

    def _setup_pdf_styles(self):
        """Define estilos visuales profesionales y sobrios (Corporate/Executive)."""
        self.styles = getSampleStyleSheet()

        # Colores Corporativos Sobrios (Dark Navy & Grey scales)
        self.color_primary = colors.HexColor("#102A43")   # Azul marino oscuro profundo
        self.color_secondary = colors.HexColor("#334E68") # Azul acero
        self.color_accent = colors.HexColor("#D32F2F")    # Rojo oscuro para alertas
        self.color_bg_light = colors.HexColor("#F0F4F8")  # Gris azulado muy claro

        # Estilo Normal Justificado
        self.styles.add(ParagraphStyle(
            name='Justify',
            parent=self.styles['Normal'],
            fontName='Helvetica',
            alignment=TA_JUSTIFY,
            leading=15,             # Interlineado más aireado
            fontSize=11,            # Letra más legible
            textColor=colors.HexColor("#243B53"),
            spaceAfter=8
        ))

        # Listas Markdown
        self.styles.add(ParagraphStyle(
            name='MarkdownList',
            parent=self.styles['Normal'],
            fontName='Helvetica',
            leftIndent=24,
            firstLineIndent=0,
            spaceAfter=6,
            bulletIndent=12,
            leading=14,
            fontSize=11,
            textColor=colors.HexColor("#243B53")
        ))

        # Bloques de Código (JSON/Code) - Estilo terminal limpio
        self.styles.add(ParagraphStyle(
            name='CodeBlock',
            fontName='Courier',
            fontSize=9,
            leading=11,
            backColor=colors.HexColor("#F5F7FA"),
            borderPadding=10,
            leftIndent=6,
            rightIndent=6,
            textColor=colors.HexColor("#102A43"),
            spaceBefore=10,
            spaceAfter=10,
            borderColor=colors.HexColor("#D9E2EC"),
            borderWidth=0.5
        ))

        # Títulos
        self.styles.add(ParagraphStyle(
            name='ReportTitle',
            parent=self.styles['Title'],
            fontSize=28,
            fontName='Helvetica-Bold',
            spaceBefore=40,
            spaceAfter=50,
            leading=36,
            textColor=self.color_primary,
            alignment=TA_CENTER
        ))

        self.styles.add(ParagraphStyle(
            name='H1',
            fontSize=18,
            fontName='Helvetica-Bold',
            spaceBefore=24,
            spaceAfter=12,
            leading=22,
            textColor=self.color_primary,
            borderPadding=0,
            borderWidth=0
        ))

        self.styles.add(ParagraphStyle(
            name='H2',
            fontSize=15,
            fontName='Helvetica-Bold',
            spaceBefore=18,
            spaceAfter=10,
            leading=20,
            textColor=self.color_secondary
        ))

        self.styles.add(ParagraphStyle(
            name='H3',
            fontSize=13,
            fontName='Helvetica-BoldOblique',
            spaceBefore=12,
            spaceAfter=6,
            leading=16,
            textColor=colors.HexColor("#486581")
        ))

        # Estilos para celdas de tabla
        self.styles.add(ParagraphStyle(
            name='CellHeader',
            fontSize=10,
            fontName='Helvetica-Bold',
            textColor=colors.white,
            alignment=TA_CENTER,
            leading=12
        ))

        self.styles.add(ParagraphStyle(
            name='CellBody',
            fontSize=10,
            fontName='Helvetica',
            textColor=colors.HexColor("#102A43"),
            alignment=TA_LEFT,
            leading=13
        ))

        # Estilo para valores de scores
        self.styles.add(ParagraphStyle(
            name='ScoreVal',
            fontSize=16,
            fontName='Helvetica-Bold',
            textColor=colors.black,
            alignment=TA_CENTER,
            leading=20
        ))

    def _get_score_color(self, score: Union[float, None]):
        if score is None: return colors.grey
        if score >= 90: return colors.HexColor("#107C10") # Verde corporativo
        if score >= 50: return colors.HexColor("#D83B01") # Naranja oscuro
        return colors.HexColor("#A80000") # Rojo oscuro

    def _extract_json_blocks(self, text: str) -> List[Dict]:
        """Extrae bloques de código JSON del texto markdown."""
        if not text: return []

        json_blocks = []

        # 1. Buscar bloques ```json
        pattern = r"```json\s*([\s\S]*?)\s*```"
        matches = re.findall(pattern, text)

        for match in matches:
            try:
                clean_json = match.strip()
                parsed = json.loads(clean_json)
                json_blocks.append(parsed)
            except json.JSONDecodeError:
                pass

        # 2. Si no hay bloques explicito, buscar bloques de codigo genericos
        if not json_blocks:
            pattern_gen = r"```\s*([\s\S]*?)\s*```"
            matches_gen = re.findall(pattern_gen, text)
            for match in matches_gen:
                try:
                    clean_json = match.strip()
                    if clean_json.startswith('{') or clean_json.startswith('['):
                        parsed = json.loads(clean_json)
                        json_blocks.append(parsed)
                except:
                    pass

        return json_blocks

    def _parse_markdown_to_flowables(self, text: str) -> List:
        """
        Parser Avanzado: Convierte Markdown a elementos PDF.
        Soporta: Headers, Listas, Bloques de Código y **TABLAS**.
        """
        if not text: return []

        flowables = []
        lines = text.split('\n')

        # Buffers para estados
        in_code = False
        code_buffer = []

        in_table = False
        table_buffer = []
        table_headers = []

        # --- Helpers ---
        def clean_and_format(raw_text):
            """Limpia XML y aplica formato inline (negritas, itálicas, código)."""
            # 1. Escapar XML básico
            safe = raw_text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')

            # 2. Negritas: **texto** -> <b>texto</b>
            safe = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', safe)

            # 3. Código inline: `texto` -> <font name="Courier">texto</font>
            safe = re.sub(r'`(.*?)`', r'<font name="Courier" backColor="#f0f0f0">\1</font>', safe)

            return safe

        def flush_table(buffer):
            """Convierte el buffer de texto de tabla en una ReportLab Table bonita."""
            if not buffer: return None

            # Procesar filas
            data = []
            for row_str in buffer:
                # Eliminar barras iniciales/finales y separar
                # Ejemplo: "| Col1 | Col2 |" -> ["Col1", "Col2"]
                cols = [c.strip() for c in row_str.strip('|').split('|')]

                # Renderizar contenido de celda como Párrafo para permitir wrapping
                formatted_cols = []
                for idx, col in enumerate(cols):
                    style = self.styles['CellHeader'] if (len(data) == 0) else self.styles['CellBody']
                    formatted_cols.append(Paragraph(clean_and_format(col), style))

                data.append(formatted_cols)

            if not data: return None

            # Calcular ancho dinámico (asumiendo ancho página LETTER - margenes 0.4 inch cada lado)
            # Ancho disponible aprox 7.7 inch
            col_count = len(data[0])
            available_width = 7.7 * inch
            col_width = available_width / col_count # Distribuir equitativamente

            t = Table(data, colWidths=[col_width] * col_count)

            # Estilo visual de la tabla
            t_style = [
                ('BACKGROUND', (0, 0), (-1, 0), self.color_primary), # Header oscuro
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                ('PADDING', (0, 0), (-1, -1), 6),
            ]

            # Filas alternas (Zebra striping)
            for i in range(1, len(data)):
                bg = self.color_bg_light if i % 2 == 0 else colors.white
                t_style.append(('BACKGROUND', (0, i), (-1, i), bg))

            t.setStyle(TableStyle(t_style))
            return t

        # --- Loop Principal ---
        for line in lines:
            stripped = line.strip()

            # 1. Detección de Bloques de Código
            if stripped.startswith("```"):
                if in_code:
                    # Fin del bloque
                    content = "\n".join(code_buffer)
                    # Escapar caracteres para preformatted
                    content = content.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                    flowables.append(Preformatted(content, self.styles["CodeBlock"]))
                    flowables.append(Spacer(1, 10))
                    code_buffer = []
                    in_code = False
                else:
                    # Inicio del bloque
                    # Si estábamos en tabla, cerrarla antes
                    if in_table:
                        t = flush_table(table_buffer)
                        if t: flowables.append(t)
                        flowables.append(Spacer(1, 10))
                        table_buffer = []
                        in_table = False

                    in_code = True
                continue

            if in_code:
                code_buffer.append(line)
                continue

            # 2. Detección de Tablas (Líneas que empiezan y terminan con |)
            # Regex básico: empieza con |, tiene contenido, termina con opcional |
            is_table_row = stripped.startswith('|') and '|' in stripped[1:]

            if is_table_row:
                # Detectar fila separadora Markdown (ej: |---|---|)
                if re.match(r'\|?[\s-]+\|[\s-]+\|?', stripped):
                    continue # Ignorar la línea separadora

                if not in_table:
                    in_table = True
                    table_buffer = []

                table_buffer.append(stripped)
                continue

            # Si no es tabla pero estábamos en una, FLUSH
            if in_table:
                t = flush_table(table_buffer)
                if t:
                    flowables.append(t)
                    flowables.append(Spacer(1, 12))
                table_buffer = []
                in_table = False

            if not stripped: continue

            # 3. Headers
            if stripped.startswith("#"):
                level = stripped.count('#')
                txt = clean_and_format(stripped.strip('# ').strip())

                if level == 1:
                    style = self.styles["H1"]
                elif level == 2:
                    style = self.styles["H2"]
                else:
                    style = self.styles["H3"]

                flowables.append(Paragraph(txt, style))

            # 4. Listas
            elif stripped.startswith("- ") or stripped.startswith("* "):
                txt = clean_and_format(stripped[2:])
                flowables.append(Paragraph(f"&bull; {txt}", self.styles["MarkdownList"]))

            # 5. Párrafo Normal
            else:
                txt = clean_and_format(stripped)
                flowables.append(Paragraph(txt, self.styles["Justify"]))

        # Limpieza final (por si el texto termina en código o tabla)
        if in_table and table_buffer:
            t = flush_table(table_buffer)
            if t: flowables.append(t)

        if in_code and code_buffer:
            content = "\n".join(code_buffer).replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            flowables.append(Preformatted(content, self.styles["CodeBlock"]))

        return flowables

    def _extract_tables_from_text(self, text: str) -> List[pd.DataFrame]:
        """Extrae tablas Markdown de un texto y devuelve una lista de DataFrames."""
        if not text: return []

        tables = []
        lines = text.split('\n')
        current_table = []
        in_table = False
        headers = None

        for line in lines:
            stripped = line.strip()
            # Detectar fila de tabla
            if stripped.startswith('|') and '|' in stripped[1:]:
                # Detectar separador |---|---|
                if re.match(r'\|?[\s-]+\|[\s-]+\|?', stripped):
                    continue

                # Limpiar celdas
                cells = [c.strip() for c in stripped.strip('|').split('|')]

                if not in_table:
                    # Nueva tabla encontrada
                    in_table = True
                    headers = cells
                else:
                    # Fila de datos
                    if headers and len(cells) != len(headers):
                        if len(cells) < len(headers):
                            cells.extend([''] * (len(headers) - len(cells)))
                        else:
                            cells = cells[:len(headers)]
                    current_table.append(cells)
            else:
                if in_table:
                    # Fin de tabla
                    if headers:
                        try:
                            df = pd.DataFrame(current_table, columns=headers)
                            if not df.empty:
                                tables.append(df)
                        except Exception as e:
                            print(f"Error creando tabla excel: {e}")

                    in_table = False
                    current_table = []
                    headers = None

        if in_table and headers:
            try:
                df = pd.DataFrame(current_table, columns=headers)
                if not df.empty:
                    tables.append(df)
            except:
                pass

        return tables

    def _write_dfs_to_sheet(self, writer, dfs: List[pd.DataFrame], sheet_name: str):
        """Escribe múltiples dataframes en una misma hoja, uno debajo del otro."""
        startrow = 0
        for i, df in enumerate(dfs):
            try:
                df.to_excel(writer, sheet_name=sheet_name, startrow=startrow, index=False)
                startrow += len(df) + 3
            except Exception as e:
                print(f"Error escribiendo tabla en hoja {sheet_name}: {e}")

    # =========================================================================
    #  METODOS DE GENERACIÓN (Integrados)
    # =========================================================================

    def generate_pdf(self) -> str:
        """Genera el reporte individual de Auditoría."""
        filename = self.base_dir / f"Reporte_SEO_{self.timestamp}.pdf"
        doc = SimpleDocTemplate(
            str(filename),
            pagesize=LETTER,
            topMargin=40,
            bottomMargin=40,
            leftMargin=0.5*inch,
            rightMargin=0.5*inch
        )
        story = []

        # --- Encabezado ---
        story.append(Paragraph("Reporte de Auditoría SEO Integral", self.styles["ReportTitle"]))

        # Meta info en una tabla invisible para alinear bien
        meta_data = [
            [Paragraph(f"<b>URL Objetivo:</b> {self.url}", self.styles["Justify"]),
             Paragraph(f"<b>ID:</b> {self.audit.id}", self.styles["Justify"])],
            [Paragraph(f"<b>Generado:</b> {datetime.now().strftime('%d/%m/%Y %H:%M')}", self.styles["Justify"]), ""]
        ]
        t_meta = Table(meta_data, colWidths=[4*inch, 3*inch])
        t_meta.setStyle(TableStyle([('VALIGN', (0,0), (-1,-1), 'TOP')]))
        story.append(t_meta)
        story.append(Spacer(1, 20))

        # --- Score Cards ---
        scores = [
            self.audit.performance_score, self.audit.seo_score,
            self.audit.accessibility_score, self.audit.best_practices_score
        ]

        # Crear celdas de colores para los scores
        def create_score_cell(label, score):
            color = self._get_score_color(score)
            return [
                Paragraph(label, self.styles["CellHeader"]),
                Paragraph(f"<font color='{color.hexval()}'>{score or 'N/A'}</font>", self.styles["ScoreVal"])
            ]

        # Tabla de scores simplificada
        data_scores = [
            ['Performance', 'SEO', 'Accessibility', 'Best Practices'],
            [
                Paragraph(f"<font color={self._get_score_color(scores[0]).hexval()}>{scores[0] or 'N/A'}</font>", self.styles["ScoreVal"]),
                Paragraph(f"<font color={self._get_score_color(scores[1]).hexval()}>{scores[1] or 'N/A'}</font>", self.styles["ScoreVal"]),
                Paragraph(f"<font color={self._get_score_color(scores[2]).hexval()}>{scores[2] or 'N/A'}</font>", self.styles["ScoreVal"]),
                Paragraph(f"<font color={self._get_score_color(scores[3]).hexval()}>{scores[3] or 'N/A'}</font>", self.styles["ScoreVal"]),
            ]
        ]

        t = Table(data_scores, colWidths=[1.8 * inch] * 4)
        t.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), self.color_primary),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.lightgrey),
            ('PADDING', (0, 0), (-1, -1), 12),
            ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
        ]))
        story.append(t)
        story.append(Spacer(1, 20))

        # --- Análisis IA ---
        story.append(Paragraph(f"Análisis de Inteligencia Artificial", self.styles["H1"]))
        analysis_text = self.ai_data.get('analysis', '')
        if analysis_text:
            story.extend(self._parse_markdown_to_flowables(analysis_text))
        else:
            story.append(Paragraph("<i>No disponible.</i>", self.styles["Justify"]))

        story.append(PageBreak())

        # --- Schemas ---
        schemas = self.seo_data.get('schema_markup', [])
        story.append(Paragraph(f"Datos Estructurados ({len(schemas)})", self.styles["H1"]))

        if schemas:
            for idx, schema in enumerate(schemas, 1):
                s_type = schema.get('@type', 'Unknown')
                story.append(Paragraph(f"{idx}. {s_type}", self.styles["H3"]))

                # Convertir JSON a string formateado
                json_str = json.dumps(schema, indent=2, ensure_ascii=False)
                # Escapado simple manual porque va a Preformatted
                json_str = json_str.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                story.append(Preformatted(json_str, self.styles["CodeBlock"]))
                story.append(Spacer(1, 10))

        doc.build(story)
        return str(filename)

    def generate_excel(self) -> str:
        filename = self.base_dir / f"Analitica_{self.timestamp}.xlsx"
        # (Tu lógica de Excel existente se mantiene igual)
        def _clean_dt(dt):
            return dt.replace(tzinfo=None) if dt and dt.tzinfo else dt

        with pd.ExcelWriter(filename, engine='openpyxl') as writer:
            summary = {
                'ID': [str(self.audit.id)], 'URL': [self.url],
                'Created At': [_clean_dt(self.audit.created_at)],
                'Performance': [self.audit.performance_score]
            }
            pd.DataFrame(summary).to_excel(writer, sheet_name='Dashboard', index=False)

            schemas = self.seo_data.get('schema_markup', [])
            if schemas:
                pd.json_normalize(schemas).to_excel(writer, sheet_name='Schemas', index=False)

        return str(filename)

    def generate_docx(self) -> str:
        """Genera el reporte en Word (DOCX)."""
        filename = self.base_dir / f"Reporte_SEO_{self.timestamp}.docx"
        doc = Document()

        # Estilos base
        self._setup_docx_styles(doc)

        # --- Encabezado ---
        title = doc.add_heading('Reporte de Auditoría SEO Integral', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Meta Info
        p = doc.add_paragraph()
        p.add_run(f"URL Objetivo: ").bold = True
        p.add_run(f"{self.url}\n")
        p.add_run(f"ID: ").bold = True
        p.add_run(f"{self.audit.id}\n")
        p.add_run(f"Generado: ").bold = True
        p.add_run(f"{datetime.now().strftime('%d/%m/%Y %H:%M')}")

        doc.add_paragraph() # Spacer

        # --- Score Cards ---
        # Tabla de 2 filas: Headers y Valores
        table = doc.add_table(rows=2, cols=4)
        table.style = 'Table Grid'

        headers = ['Performance', 'SEO', 'Accessibility', 'Best Practices']
        scores = [
            self.audit.performance_score, self.audit.seo_score,
            self.audit.accessibility_score, self.audit.best_practices_score
        ]

        # Llenar headers
        for i, header in enumerate(headers):
            cell = table.cell(0, i)
            cell.text = header
            # Centrar y negrita
            tc = cell.paragraphs[0]
            tc.alignment = WD_ALIGN_PARAGRAPH.CENTER
            tc.runs[0].bold = True
            self._set_cell_background(cell, "102A43") # Azul oscuro
            tc.runs[0].font.color.rgb = RGBColor(255, 255, 255)

        # Llenar valores
        for i, score in enumerate(scores):
            cell = table.cell(1, i)
            cell.text = str(score or 'N/A')
            tc = cell.paragraphs[0]
            tc.alignment = WD_ALIGN_PARAGRAPH.CENTER
            tc.runs[0].font.size = Pt(16)
            tc.runs[0].bold = True

            # Color segun score
            color = self._get_score_color_docx(score)
            tc.runs[0].font.color.rgb = color

        doc.add_paragraph()

        # --- Análisis IA ---
        doc.add_heading('Análisis de Inteligencia Artificial', level=1)
        analysis_text = self.ai_data.get('analysis', '')
        if analysis_text:
            self._parse_markdown_to_docx(doc, analysis_text)
        else:
            doc.add_paragraph("No disponible.", style='Normal')

        doc.add_page_break()

        # --- Schemas ---
        schemas = self.seo_data.get('schema_markup', [])
        doc.add_heading(f"Datos Estructurados ({len(schemas)})", level=1)

        if schemas:
            for idx, schema in enumerate(schemas, 1):
                s_type = schema.get('@type', 'Unknown')
                doc.add_heading(f"{idx}. {s_type}", level=3)

                json_str = json.dumps(schema, indent=2, ensure_ascii=False)
                # Code block simulado
                p = doc.add_paragraph(json_str)
                p.style = 'Quote' # Usar estilo Quote o crear uno custom de código
                # Aplicar fuente monoespaciada
                for run in p.runs:
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)

        doc.save(filename)
        return str(filename)

    # =========================================================================
    #  COMPARATIVOS (Benchmarking)
    # =========================================================================

    def generate_comparison_reports(self, comparison_data: Union[Dict, Any]) -> Dict[str, str]:
        # 1. Normalización
        if hasattr(comparison_data, 'model_dump'):
            data = comparison_data.model_dump()
        elif hasattr(comparison_data, 'dict'):
            data = comparison_data.dict()
        else:
            data = comparison_data

        ts = datetime.now().strftime("%Y%m%d_%H%M")
        pdf_path = self.base_dir / f"Benchmark_Report_{ts}.pdf"
        xlsx_path = self.base_dir / f"Benchmark_Data_{ts}.xlsx"
        word_path = self.base_dir / f"Benchmark_Report_{ts}.docx"

        self._create_comparison_pdf(data, pdf_path)
        self._create_comparison_excel(data, xlsx_path)
        self._create_comparison_docx(data, word_path)

        return {"pdf_path": str(pdf_path), "xlsx_path": str(xlsx_path), "word_path": str(word_path)}

    def _create_comparison_docx(self, data: dict, filename: Path):
        """Genera el reporte comparativo en Word (DOCX)."""
        doc = Document()

        # Estilos base
        self._setup_docx_styles(doc)

        # --- Encabezado ---
        title = doc.add_heading('Reporte Comparativo de Auditoría SEO', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Meta Info - Suponiendo que hay un campo 'url' y 'id' en los datos de comparación
        for comp in data.get('comparisons', []):
            url = comp.get('url', 'Desconocida')
            audit_id = comp.get('id', 'N/A')

            p = doc.add_paragraph()
            p.add_run(f"URL: ").bold = True
            p.add_run(f"{url}\n")
            p.add_run(f"ID: ").bold = True
            p.add_run(f"{audit_id}\n")
            p.add_run(f"Generado: ").bold = True
            p.add_run(f"{datetime.now().strftime('%d/%m/%Y %H:%M')}")
            doc.add_paragraph() # Spacer

        # --- Tabla Comparativa ---
        # Suponiendo que hay una sección 'table' en los datos de comparación
        if 'table' in data:
            table_data = data['table']
            # Extraer headers asumiendo que son la primera fila
            headers = table_data[0] if table_data else []
            # Crear tabla
            table = doc.add_table(rows=0, cols=len(headers))
            table.style = 'Table Grid'
            table.autofit = True

            # Agregar encabezados
            hdr_cells = table.add_row().cells
            for i, header in enumerate(headers):
                hdr_cells[i].text = str(header)
                # Estilo header
                self._set_cell_background(hdr_cells[i], "102A43") # Azul oscuro
                for layout in hdr_cells[i].paragraphs:
                    for run in layout.runs:
                        run.font.color.rgb = RGBColor(255, 255, 255)
                        run.bold = True

            # Agregar datos
            for row in table_data[1:]:
                if not row: continue # Saltar filas vacías
                row_cells = table.add_row().cells
                for i, cell_value in enumerate(row):
                    if i < len(row_cells):
                        p = row_cells[i].paragraphs[0]
                        # Limpiar y agregar texto
                        p.clear()
                        p.add_run(str(cell_value)).font.size = Pt(10.5)

        doc.save(filename)

    def _setup_docx_styles(self, doc):
        """Configura estilos del documento Word."""
        try:
            # Puedes personalizar estilos existentes o crear nuevos
            # Estilo Normal
            styles = doc.styles
            if 'Normal' in styles:
                style = styles['Normal']
                font = style.font
                font.name = 'Helvetica'
                font.size = Pt(11)
                font.color.rgb = RGBColor(36, 59, 83) # #243B53
        except:
            pass

    def _get_score_color_docx(self, score):
        if score is None: return RGBColor(128, 128, 128)
        if score >= 90: return RGBColor(16, 124, 16) # #107C10
        if score >= 50: return RGBColor(216, 59, 1) # #D83B01
        return RGBColor(168, 0, 0) # #A80000

    def _set_cell_background(self, cell, hex_color):
        """Helper para poner color de fondo a una celda."""
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:val'), 'clear')
        shading_elm.set(qn('w:color'), 'auto')
        shading_elm.set(qn('w:fill'), hex_color)
        cell._tc.get_or_add_tcPr().append(shading_elm)

    def _parse_markdown_to_docx(self, doc, text: str):
        """Parsea Markdown básico a elementos de Word."""
        lines = text.split('\n')
        in_code = False
        code_buffer = []
        in_table = False
        table_buffer = []

        # Helper para procesar formato inline (**negrita**)
        def add_formatted_text(paragraph, text):
            # Split por **
            parts = re.split(r'(\*\*.*?\*\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = paragraph.add_run(part[2:-2])
                    run.bold = True
                else:
                    # Buscar `code` inline
                    subparts = re.split(r'(`.*?`)', part)
                    for subpart in subparts:
                        if subpart.startswith('`') and subpart.endswith('`'):
                            run = paragraph.add_run(subpart[1:-1])
                            run.font.name = 'Courier New'
                            run.font.highlight_color = 16 # Gris claro (WD_COLOR_INDEX.GRAY_25 aprox)
                        else:
                            paragraph.add_run(subpart)

        def flush_table_docx(buffer):
            if not buffer: return
            # Determinar dimensiones
            rows_data = []
            for line in buffer:
                # Basic pipe parsing
                cols = [c.strip() for c in line.strip('|').split('|')]
                rows_data.append(cols)

            if not rows_data: return

            # Crear tabla
            num_cols = len(rows_data[0])
            table = doc.add_table(rows=0, cols=num_cols)
            table.style = 'Table Grid'
            table.autofit = True

            # Header
            hdr_cells = table.add_row().cells
            for i, col_txt in enumerate(rows_data[0]):
                if i < len(hdr_cells):
                    hdr_cells[i].text = col_txt
                    # Estilo header
                    self._set_cell_background(hdr_cells[i], "102A43") # Azul oscuro
                    for layout in hdr_cells[i].paragraphs:
                        for run in layout.runs:
                            run.font.color.rgb = RGBColor(255, 255, 255)
                            run.bold = True

            # Data
            for row_cols in rows_data[1:]:
                # Check if separator row
                if set(''.join(row_cols)) <= set(['-', ' ', ':']): continue

                row_cells = table.add_row().cells
                for i, col_txt in enumerate(row_cols):
                    if i < len(row_cells):
                        p = row_cells[i].paragraphs[0]
                        add_formatted_text(p, col_txt)

            doc.add_paragraph() # Spacer

        for line in lines:
            stripped = line.strip()

            # Code Blocks
            if stripped.startswith("```"):
                if in_code:
                    # Flush code
                    p = doc.add_paragraph('\n'.join(code_buffer))
                    p.style = 'Quote' # Estilo simple para código
                    for run in p.runs:
                        run.font.name = 'Courier New'
                        run.font.size = Pt(9)
                    code_buffer = []
                    in_code = False
                else:
                    if in_table:
                        flush_table_docx(table_buffer)
                        table_buffer = []
                        in_table = False
                    in_code = True
                continue

            if in_code:
                code_buffer.append(line)
                continue

            # Tables
            if stripped.startswith('|') and '|' in stripped[1:]:
                if not in_table:
                    in_table = True
                    table_buffer = []
                table_buffer.append(stripped)
                continue

            if in_table:
                flush_table_docx(table_buffer)
                table_buffer = []
                in_table = False

            if not stripped: continue

            # Headings
            if stripped.startswith('#'):
                level = stripped.count('#')
                txt = stripped.strip('# ').strip()
                # Word supports levels 1-9
                lvl = min(level, 9)
                doc.add_heading(txt, level=lvl)

            # Lists
            elif stripped.startswith('- ') or stripped.startswith('* '):
                txt = stripped[2:]
                p = doc.add_paragraph(style='List Bullet')
                add_formatted_text(p, txt)

            else:
                p = doc.add_paragraph()
                add_formatted_text(p, stripped)

        # Final flush
        if in_table: flush_table_docx(table_buffer)
        if in_code and code_buffer:
            p = doc.add_paragraph('\n'.join(code_buffer))
            p.style = 'Quote'
            for run in p.runs:
                run.font.name = 'Courier New'

    def generate_all(self):
        return {
            "pdf_path": self.generate_pdf(),
            "xlsx_path": self.generate_excel(),
            "word_path": self.generate_docx()
        }

